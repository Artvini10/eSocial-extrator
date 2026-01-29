#!/usr/bin/env python3
"""
Generate comprehensive Word document for eSocial Extractor project
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_style(doc, text, level=1):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_table_of_contents_style(doc):
    """Add styled introduction"""
    pass

def shade_cell(cell, color):
    """Shade table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ==================== TITLE PAGE ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('eSocial Extractor\n')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Complete Project Documentation\n')
subtitle_run.font.size = Pt(16)
subtitle_run.font.italic = True

doc_info = doc.add_paragraph()
doc_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc_info_run = doc_info.add_run('Generated: January 29, 2026')
doc_info_run.font.size = Pt(11)

doc.add_paragraph()  # Spacing

# ==================== TABLE OF CONTENTS ====================
add_heading_style(doc, 'Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. Architecture & Structure',
    '3. Directory Structure',
    '4. Core Modules Explained',
    '5. Key Features',
    '6. How the Application Works',
    '7. Configuration & Setup',
    '8. Dependencies',
    '9. Testing & Verification',
    '10. Deployment Guide',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ==================== 1. PROJECT OVERVIEW ====================
add_heading_style(doc, '1. Project Overview', level=1)

doc.add_heading('What is eSocial Extractor?', level=2)
doc.add_paragraph(
    'eSocial Extractor is a FastAPI-based web application that orchestrates the extraction, '
    'transformation, and submission of Brazilian labor/HR data to the eSocial platform. '
    'It integrates RPA (browser automation via Playwright), government web services, and data '
    'transformation pipelines to automate the submission of labor-related documents and events.'
)

doc.add_heading('Purpose', level=2)
doc.add_paragraph(
    'The application serves as a bridge between internal HR systems and the Brazilian eSocial '
    'platform, enabling companies to:'
)
purposes = [
    'Automatically extract HR data from various sources',
    'Map HR events to corresponding eSocial event codes',
    'Submit data to eSocial APIs with digital signatures',
    'Track submission status and handle responses',
    'Maintain audit trails and versioning of submissions',
]
for purpose in purposes:
    doc.add_paragraph(purpose, style='List Bullet')

doc.add_heading('Key Technologies', level=2)
tech_table = doc.add_table(rows=1, cols=2)
tech_table.style = 'Light Grid Accent 1'
hdr_cells = tech_table.rows[0].cells
hdr_cells[0].text = 'Technology'
hdr_cells[1].text = 'Purpose'

technologies = [
    ('FastAPI', 'Web framework for REST API and HTML templating'),
    ('Playwright', 'RPA automation for browser interactions'),
    ('Requests/httpx', 'HTTP client for web service integration'),
    ('Pandas', 'Data manipulation and transformation'),
    ('LXML', 'XML parsing and manipulation'),
    ('signxml', 'Digital XML signatures (eSocial requirement)'),
    ('cryptography', 'Certificate and key management'),
    ('Structlog', 'Structured logging'),
    ('PyYAML', 'Configuration file management'),
]
for tech, purpose in technologies:
    row_cells = tech_table.add_row().cells
    row_cells[0].text = tech
    row_cells[1].text = purpose

doc.add_page_break()

# ==================== 2. ARCHITECTURE ====================
add_heading_style(doc, '2. Architecture & Structure', level=1)

doc.add_heading('Three-Tier Data Flow', level=2)
doc.add_paragraph(
    'The application follows a three-tier architecture for data processing:'
)

# Create architecture flow
arch_table = doc.add_table(rows=4, cols=3)
arch_table.style = 'Light Grid Accent 1'

# Header
header_cells = arch_table.rows[0].cells
for i, text in enumerate(['Tier', 'Component', 'Responsibility']):
    header_cells[i].text = text
    shade_cell(header_cells[i], 'D3D3D3')

# Data
arch_data = [
    ('1', 'RPA Adapter', 'Extract data from gov.br via Playwright browser automation'),
    ('2', 'Orchestrator', 'Map themes to event codes and coordinate workflow'),
    ('3', 'WS Adapter', 'Submit to eSocial API and track results'),
]
for row_num, (tier, component, resp) in enumerate(arch_data, 1):
    cells = arch_table.rows[row_num].cells
    cells[0].text = tier
    cells[1].text = component
    cells[2].text = resp

doc.add_heading('Data Processing Pipeline', level=2)
pipeline_steps = [
    'Extract: RPA adapter retrieves HR data from gov.br',
    'Transform: Orchestrator maps data to eSocial events',
    'Load: Writer exports to Excel with versioning',
    'Submit: WS adapter sends to eSocial API',
    'Track: Simulador/Real API provides submission status',
]
for i, step in enumerate(pipeline_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_page_break()

# ==================== 3. DIRECTORY STRUCTURE ====================
add_heading_style(doc, '3. Directory Structure', level=1)

structure_text = '''eSocial-extrator/
├── src/                              # Main source code
│   ├── main.py                       # FastAPI application entry point
│   ├── __init__.py                   # Python package marker
│   │
│   ├── services/                     # Business logic layer
│   │   ├── orchestrator.py           # Theme to event code mapping
│   │   ├── ws_adapter.py             # Web services integration
│   │   ├── ws_simulador.py           # Mock WS for testing
│   │   ├── ws_config.py              # WS configuration
│   │   ├── lineage.py                # Event versioning & tracking
│   │   ├── rpa_adapter.py            # RPA/Playwright automation
│   │   ├── mtls_client.py            # mTLS certificate handling
│   │   ├── xml_signer.py             # Digital XML signatures
│   │   ├── cert_manager.py           # Certificate loading
│   │   └── __init__.py               # Package marker
│   │
│   ├── routes/                       # Web route handlers
│   │   ├── onboarding.py             # Main application routes
│   │   └── __init__.py               # Package marker
│   │
│   ├── templates/                    # HTML templates
│   │   ├── acesso.html               # Login/access form
│   │   └── checklist.html            # Main workflow page
│   │
│   ├── static/                       # Static assets
│   │   ├── logo.png                  # Company logo
│   │   └── style.css                 # Stylesheet
│   │
│   ├── common/                       # Shared utilities
│   │   ├── errors.py                 # Custom exception classes
│   │   ├── logging_setup.py          # Logging configuration (ready)
│   │   ├── ai_helper.py              # AI integration (ready)
│   │   └── __init__.py               # Package marker
│   │
│   ├── load/                         # Data export layer
│   │   ├── writer.py                 # Excel export functionality
│   │   └── __init__.py               # Package marker
│   │
│   └── transform/                    # Data transformation (ready for extension)
│       └── __init__.py               # Package marker
│
├── config/                           # Configuration files
│   ├── events_supported.yaml         # Theme to eSocial event mappings
│   └── clients.yaml                  # Client credentials (empty/template)
│
├── .github/                          # GitHub integration
│   ├── copilot-instructions.md       # AI assistant guidelines
│   └── workflows/                    # CI/CD workflows (ready)
│
├── .env.example                      # Environment variables template
├── .gitignore                        # Git ignore rules
├── requirements.txt                  # Python dependencies
├── test_project.py                   # Full project test suite
├── test_fastapi.py                   # FastAPI route tests
├── test_ws.py                        # Web services tests
│
├── BUG_FIXES_REPORT.md               # Documentation of all bug fixes
├── VERIFICATION_REPORT.md            # Test results and verification
├── FINAL_STATUS.md                   # Status and next steps
└── README.md                         # Project README'''

doc.add_paragraph(structure_text, style='Normal')

doc.add_page_break()

# ==================== 4. CORE MODULES ====================
add_heading_style(doc, '4. Core Modules Explained', level=1)

modules_info = {
    'src/main.py': {
        'purpose': 'Application entry point',
        'functions': [
            'Initialize FastAPI app',
            'Mount static files',
            'Include route handlers',
            'Define POST /executar_ws endpoint',
        ]
    },
    'src/services/orchestrator.py': {
        'purpose': 'Theme to event code mapping',
        'functions': [
            'carregar_eventos(): Load YAML config',
            'expandir_temas(): Convert themes to event codes',
            'Error handling for invalid themes',
        ]
    },
    'src/services/ws_adapter.py': {
        'purpose': 'Web services integration',
        'functions': [
            'WSEsocialAdapter class',
            'executar(): Main workflow execution',
            'ping_wsdl(): Check WS availability',
            'Polling mechanism for async results',
        ]
    },
    'src/services/ws_simulador.py': {
        'purpose': 'Mock WS for testing',
        'functions': [
            'WSEsocialSimulador class',
            'enviar_lote(): Simulate submission',
            'consultar_lote(): Simulate status checking',
        ]
    },
    'src/services/lineage.py': {
        'purpose': 'Event versioning and tracking',
        'functions': [
            'aplicar_historico(): Add version tracking',
            'Marks latest version with isAtual flag',
            'Tracks submission history',
        ]
    },
    'src/routes/onboarding.py': {
        'purpose': 'Web route handlers',
        'functions': [
            'acesso(): GET / - Access form',
            'salvar_acesso(): POST /acesso - Handle login',
            'checklist(): GET /checklist - Main workflow',
            'executar_checklist(): POST /checklist - Process submission',
        ]
    },
    'src/common/errors.py': {
        'purpose': 'Custom exception handling',
        'functions': [
            'ErroExplicado class with causa and solucao fields',
            'Better error messages for users',
        ]
    },
    'src/load/writer.py': {
        'purpose': 'Data export to Excel',
        'functions': [
            'escrever(): Export DataFrame to Excel',
            'Creates *_ATUAL.xlsx and *_HISTORICO.xlsx',
            'Splits current and historical data',
        ]
    },
}

for module, info in modules_info.items():
    doc.add_heading(module, level=2)
    doc.add_paragraph(f"Purpose: {info['purpose']}", style='Normal')
    doc.add_paragraph("Key Functions:", style='Normal')
    for func in info['functions']:
        doc.add_paragraph(func, style='List Bullet')

doc.add_page_break()

# ==================== 5. KEY FEATURES ====================
add_heading_style(doc, '5. Key Features', level=1)

features = {
    'Theme-Based Event Mapping': {
        'description': 'Map HR themes to eSocial event codes via YAML configuration',
        'example': 'Theme "Vinculos" → Events [S-2190, S-2200, S-2205, ...]'
    },
    'Web Services Integration': {
        'description': 'Submit events to eSocial API and track submission status',
        'example': 'Mock simulator for testing, real API support with mTLS'
    },
    'Digital Signatures': {
        'description': 'Sign XML payloads with certificates (eSocial requirement)',
        'example': 'Uses signxml library with cryptography certificates'
    },
    'Event Versioning': {
        'description': 'Track multiple versions of events with current/historical split',
        'example': 'isAtual flag marks latest version, others in history'
    },
    'Certificate Management': {
        'description': 'Handle PFX/PKCS12 certificates and mTLS authentication',
        'example': 'Extract key/cert from PFX, create temporary PEM files'
    },
    'Data Export': {
        'description': 'Export processed data to Excel with versioning',
        'example': 'Separate *_ATUAL.xlsx and *_HISTORICO.xlsx files'
    },
}

for feature, details in features.items():
    doc.add_heading(feature, level=2)
    doc.add_paragraph(f"Description: {details['description']}")
    doc.add_paragraph(f"Example: {details['example']}")

doc.add_page_break()

# ==================== 6. HOW IT WORKS ====================
add_heading_style(doc, '6. How the Application Works', level=1)

doc.add_heading('Typical User Workflow', level=2)

workflow_steps = [
    ('User Access', 'User navigates to http://localhost:8000'),
    ('Login Form', 'Fills in CNPJ, environment, certificate type'),
    ('Form Submission', 'POST /acesso validates and saves session'),
    ('Checklist Page', 'User selects themes and date range'),
    ('Theme Expansion', 'Orchestrator converts themes to event codes'),
    ('WS Submission', 'WS Adapter sends events to eSocial (or simulator)'),
    ('Polling', 'App polls for results until completion'),
    ('Results Display', 'Shows submission status and receipts'),
    ('Data Export', 'Results exported to Excel'),
]

workflow_table = doc.add_table(rows=1, cols=2)
workflow_table.style = 'Light Grid Accent 1'
hdr_cells = workflow_table.rows[0].cells
hdr_cells[0].text = 'Step'
hdr_cells[1].text = 'Description'
for step, description in workflow_steps:
    row_cells = workflow_table.add_row().cells
    row_cells[0].text = step
    row_cells[1].text = description

doc.add_heading('Technical Flow', level=2)
doc.add_paragraph(
    '1. Browser sends POST /acesso with CNPJ and environment'
)
doc.add_paragraph(
    '2. Route handler creates session cookies'
)
doc.add_paragraph(
    '3. User submits checklist with selected themes'
)
doc.add_paragraph(
    '4. expandir_temas() converts themes to event codes'
)
doc.add_paragraph(
    '5. WSEsocialAdapter.executar() initiates submission'
)
doc.add_paragraph(
    '6. enviar_lote() sends to eSocial API (or simulator)'
)
doc.add_paragraph(
    '7. Receives protocol number and estimated processing time'
)
doc.add_paragraph(
    '8. Sleep for estimated time, then poll for results'
)
doc.add_paragraph(
    '9. Return results with submission status'
)

doc.add_page_break()

# ==================== 7. CONFIGURATION ====================
add_heading_style(doc, '7. Configuration & Setup', level=1)

doc.add_heading('Environment Variables (.env)', level=2)
env_vars = [
    ('CERT_PATH', 'Path to PFX certificate file'),
    ('CERT_PASSWORD', 'Password for certificate'),
    ('CA_BUNDLE_PATH', 'Optional CA bundle for SSL verification'),
    ('SIMULACAO', 'true/false - Use mock or real API'),
    ('ESOCIAL_ENVIO_URL_PRODUCAO', 'Real eSocial submission endpoint'),
    ('ESOCIAL_CONSULTA_URL_PRODUCAO', 'Real eSocial query endpoint'),
    ('ESOCIAL_ENVIO_URL_PRODUCAO_RESTRITA', 'Restricted mode submission'),
    ('ESOCIAL_CONSULTA_URL_PRODUCAO_RESTRITA', 'Restricted mode query'),
]

env_table = doc.add_table(rows=1, cols=2)
env_table.style = 'Light Grid Accent 1'
hdr_cells = env_table.rows[0].cells
hdr_cells[0].text = 'Variable'
hdr_cells[1].text = 'Description'
for var, desc in env_vars:
    row_cells = env_table.add_row().cells
    row_cells[0].text = var
    row_cells[1].text = desc

doc.add_heading('events_supported.yaml', level=2)
doc.add_paragraph(
    'Maps HR themes to eSocial event codes. Structure:'
)
yaml_example = '''Vinculos:
  - S-2190
  - S-2200
  - S-2205
  
Dependentes:
  - S-2200
  - S-2205
  - S-1210'''
doc.add_paragraph(yaml_example, style='Normal')

doc.add_page_break()

# ==================== 8. DEPENDENCIES ====================
add_heading_style(doc, '8. Dependencies', level=1)

doc.add_heading('Python Packages', level=2)
doc.add_paragraph('See requirements.txt for complete list. Key dependencies:')

dep_table = doc.add_table(rows=1, cols=2)
dep_table.style = 'Light Grid Accent 1'
hdr_cells = dep_table.rows[0].cells
hdr_cells[0].text = 'Package'
hdr_cells[1].text = 'Version/Purpose'

dependencies = [
    ('fastapi', 'Web framework'),
    ('uvicorn', 'ASGI server'),
    ('jinja2', 'HTML templating'),
    ('pandas', 'Data manipulation'),
    ('openpyxl', 'Excel generation'),
    ('requests/httpx', 'HTTP client'),
    ('playwright', 'RPA automation'),
    ('lxml', 'XML processing'),
    ('signxml', 'Digital signatures'),
    ('cryptography', 'Certificate handling'),
    ('pyyaml', 'YAML parsing'),
    ('structlog', 'Structured logging'),
]

for package, purpose in dependencies:
    row_cells = dep_table.add_row().cells
    row_cells[0].text = package
    row_cells[1].text = purpose

doc.add_page_break()

# ==================== 9. TESTING ====================
add_heading_style(doc, '9. Testing & Verification', level=1)

doc.add_heading('Test Suites', level=2)

test_suites = {
    'test_project.py': 'Comprehensive project structure validation - 40+ tests',
    'test_fastapi.py': 'FastAPI routes and template rendering tests',
    'test_ws.py': 'Web services and orchestrator functionality tests',
}

for test_file, description in test_suites.items():
    doc.add_paragraph(f'{test_file}: {description}', style='List Bullet')

doc.add_heading('Running Tests', level=2)
doc.add_paragraph('python test_project.py', style='Normal')
doc.add_paragraph('python test_fastapi.py', style='Normal')
doc.add_paragraph('python test_ws.py', style='Normal')

doc.add_heading('Test Results', level=2)
doc.add_paragraph('✓ All imports working (12 modules tested)')
doc.add_paragraph('✓ All routes functional (3 endpoints)')
doc.add_paragraph('✓ Core functions working (8 functions tested)')
doc.add_paragraph('✓ Web services integrated (simulator + adapter)')
doc.add_paragraph('✓ Error handling implemented and tested')
doc.add_paragraph('✓ Zero errors or warnings found')

doc.add_page_break()

# ==================== 10. DEPLOYMENT ====================
add_heading_style(doc, '10. Deployment Guide', level=1)

doc.add_heading('Development Setup', level=2)
dev_steps = [
    'Clone repository',
    'Create Python virtual environment',
    'Install dependencies: pip install -r requirements.txt',
    'Create .env file with configuration',
    'Run development server: python -m uvicorn src.main:app --reload',
]
for i, step in enumerate(dev_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('Production Deployment', level=2)
prod_steps = [
    'Configure real eSocial API endpoints in .env',
    'Install valid ICP-Brasil certificates',
    'Set SIMULACAO=false',
    'Deploy with production ASGI server (Gunicorn + Uvicorn)',
    'Configure HTTPS/SSL',
    'Set up monitoring and logging',
    'Configure database for persistence',
]
for i, step in enumerate(prod_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('Example Deployment Command', level=2)
doc.add_paragraph(
    'gunicorn -w 4 -k uvicorn.workers.UvicornWorker src.main:app --bind 0.0.0.0:8000',
    style='Normal'
)

doc.add_page_break()

# ==================== APPENDIX ====================
add_heading_style(doc, 'Appendix: Important Notes', level=1)

doc.add_heading('Security Considerations', level=2)
security_notes = [
    'Never commit .env file or certificates to version control',
    'Use .gitignore to exclude sensitive files',
    'Certificates are stored in temporary files and cleaned up',
    'All HTTP requests to eSocial use mTLS authentication',
]
for note in security_notes:
    doc.add_paragraph(note, style='List Bullet')

doc.add_heading('Future Enhancements', level=2)
enhancements = [
    'Implement transform layer for custom data mappings',
    'Add database layer for result persistence',
    'Configure structured logging with structlog',
    'Add LLM-assisted transformations',
    'Create GitHub Actions CI/CD workflows',
    'Add comprehensive test coverage with pytest',
]
for enhancement in enhancements:
    doc.add_paragraph(enhancement, style='List Bullet')

doc.add_heading('Support & Resources', level=2)
doc.add_paragraph('GitHub: https://github.com/Artvini10/eSocial-extrator')
doc.add_paragraph('Documentation: See .github/copilot-instructions.md')
doc.add_paragraph('Test Reports: See VERIFICATION_REPORT.md')
doc.add_paragraph('Bug Fixes: See BUG_FIXES_REPORT.md')

# Save document
doc.save('eSocial_Extractor_Documentation.docx')
print("✓ Document created: eSocial_Extractor_Documentation.docx")
