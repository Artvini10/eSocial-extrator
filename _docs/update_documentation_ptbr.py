#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Atualizar Documentação em Português (.docx)
Este script atualiza o arquivo eSocial_Extractor_Documentacao_PTBR.docx existente
com a estrutura de código atual e detalhes de implementação.

Uso: python update_documentation_ptbr.py
"""

import sys
import os
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

if sys.stdout.encoding != 'utf-8':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')


def shade_cell(cell, fill_color):
    """Aplicar sombreamento a uma célula de tabela."""
    cell_xml = cell._element
    cell_pr = cell_xml.get_or_add_tcPr()
    shade_obj = cell_pr.first_child_found_in("w:shd")
    if shade_obj is None:
        from docx.oxml import parse_xml
        shade_obj = parse_xml(f'<w:shd {{"w:fill":"{fill_color}"}}>')
        cell_pr.append(shade_obj)
    else:
        shade_obj.set('w:fill', fill_color)


def get_code_structure():
    """Obter estrutura de código atual e estatísticas."""
    src_dir = Path("src")
    stats = {
        "modulos": [],
        "linhas_total": 0,
        "arquivos_total": 0
    }
    
    for py_file in src_dir.rglob("*.py"):
        if "__pycache__" in str(py_file):
            continue
        try:
            with open(py_file, 'r', encoding='utf-8', errors='ignore') as f:
                linhas = len(f.readlines())
                stats["linhas_total"] += linhas
                stats["arquivos_total"] += 1
                rel_path = py_file.relative_to(src_dir)
                stats["modulos"].append({
                    "caminho": str(rel_path),
                    "linhas": linhas
                })
        except:
            pass
    
    return stats


def get_git_status():
    """Obter status git atual."""
    try:
        result = subprocess.run(
            ["git", "log", "-1", "--pretty=format:%h %s (%ar)"],
            capture_output=True,
            text=True,
            timeout=5
        )
        return result.stdout.strip() if result.returncode == 0 else "N/A"
    except:
        return "N/A"


def update_documentation():
    """Atualizar a documentação em português com estrutura de código atual."""
    
    doc_path = "eSocial_Extractor_Documentacao_PTBR.docx"
    
    if not os.path.exists(doc_path):
        print(f"[ERROR] Documento nao encontrado: {doc_path}")
        return False
    
    try:
        doc = Document(doc_path)
        
        # Obter estatísticas atuais
        stats = get_code_structure()
        git_status = get_git_status()
        update_time = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        
        # Encontrar e atualizar seção de estatísticas
        updated_stats = False
        for i, para in enumerate(doc.paragraphs):
            if "Estatísticas do Projeto" in para.text or "Status do Código" in para.text:
                # Encontrar próximo parágrafo e atualizar com estatísticas atuais
                if i + 2 < len(doc.paragraphs):
                    summary_para = doc.paragraphs[i + 1]
                    summary_para.text = f"Total de Arquivos: {stats['arquivos_total']} | Total de Linhas: {stats['linhas_total']} | Atualizado: {update_time}"
                    updated_stats = True
                    break
        
        # Atualizar rodapé de versão/status se existir
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Última Atualização" in cell.text or "Atualizado" in cell.text:
                        cell.text = f"Última Atualização: {update_time}\nStatus Git: {git_status}"
        
        # Adicionar nota de atualização ao final se ainda não existir
        last_para = doc.paragraphs[-1]
        if "Última Atualização" not in last_para.text:
            doc.add_paragraph(f"\n📋 Última Atualização: {update_time}\n🔗 Commit Mais Recente: {git_status}")
        
        # Salvar documento atualizado
        doc.save(doc_path)
        print(f"[OK] Documentacao em portugues atualizada: {doc_path}")
        print(f"     Arquivos: {stats['arquivos_total']} | Linhas: {stats['linhas_total']}")
        print(f"     Atualizado em: {update_time}")
        print(f"     Commit: {git_status}")
        return True
        
    except Exception as e:
        print(f"[ERROR] Erro ao atualizar documentacao: {e}")
        return False


def main():
    """Ponto de entrada principal."""
    print("\n" + "="*60)
    print("[PT-BR] Atualizando Documentacao em Portugues")
    print("="*60)
    
    success = update_documentation()
    
    if success:
        print("\n[OK] Atualizacao de documentacao concluida com sucesso!")
    else:
        print("\n[ERROR] Falha na atualizacao de documentacao!")
    
    print("="*60 + "\n")


if __name__ == "__main__":
    main()
