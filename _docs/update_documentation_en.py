#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Update English Documentation (.docx)
This script updates the existing eSocial_Extractor_Documentation.docx file
with current code structure and implementation details.

Usage: python update_documentation_en.py
"""

import sys
import os
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

if sys.stdout.encoding != 'utf-8':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')


def shade_cell(cell, fill_color):
    """Apply shading to a table cell."""
    cell_xml = cell._element
    cell_pr = cell_xml.get_or_add_tcPr()
    shade_obj = cell_pr.first_child_found_in("w:shd")
    if shade_obj is None:
        from docx.oxml import parse_xml
        shade_obj = parse_xml(f'<w:shd {{"w:fill":"{fill_color}"}}>')
        cell_pr.append(shade_obj)
    else:
        shade_obj.set('w:fill', fill_color)


def get_code_structure():
    """Get current code structure and statistics."""
    src_dir = Path("src")
    stats = {
        "modules": [],
        "line_count": 0,
        "file_count": 0
    }
    
    for py_file in src_dir.rglob("*.py"):
        if "__pycache__" in str(py_file):
            continue
        try:
            with open(py_file, 'r', encoding='utf-8', errors='ignore') as f:
                lines = len(f.readlines())
                stats["line_count"] += lines
                stats["file_count"] += 1
                rel_path = py_file.relative_to(src_dir)
                stats["modules"].append({
                    "path": str(rel_path),
                    "lines": lines
                })
        except:
            pass
    
    return stats


def get_git_status():
    """Get current git status."""
    try:
        result = subprocess.run(
            ["git", "log", "-1", "--pretty=format:%h %s (%ar)"],
            capture_output=True,
            text=True,
            timeout=5
        )
        return result.stdout.strip() if result.returncode == 0 else "N/A"
    except:
        return "N/A"


def update_documentation():
    """Update the English documentation with current code structure."""
    
    doc_path = Path("..") / "Documentação" / "eSocial_Extractor_Documentation.docx"
    
    if not doc_path.exists():
        print(f"[ERROR] Document not found: {doc_path}")
        return False
    
    try:
        doc = Document(doc_path)
        
        # Get current statistics
        stats = get_code_structure()
        git_status = get_git_status()
        update_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Find and update "Project Statistics" or "Codebase Status" section
        updated_stats = False
        for i, para in enumerate(doc.paragraphs):
            if "Project Statistics" in para.text or "Codebase Status" in para.text:
                # Find the next paragraph and update it with current stats
                if i + 2 < len(doc.paragraphs):
                    summary_para = doc.paragraphs[i + 1]
                    summary_para.text = f"Total Files: {stats['file_count']} | Total Lines: {stats['line_count']} | Updated: {update_time}"
                    updated_stats = True
                    break
        
        # Update the version/status footer if it exists
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Last Updated" in cell.text or "Updated" in cell.text:
                        cell.text = f"Last Updated: {update_time}\nGit Status: {git_status}"
        
        # Add an update note at the end if not already present
        last_para = doc.paragraphs[-1]
        if "Last Updated" not in last_para.text:
            doc.add_paragraph(f"\n📋 Last Updated: {update_time}\n🔗 Latest Commit: {git_status}")
        
        # Save the updated document
        doc.save(doc_path)
        print(f"[OK] English documentation updated: {doc_path}")
        print(f"     Files: {stats['file_count']} | Lines: {stats['line_count']}")
        print(f"     Updated at: {update_time}")
        print(f"     Commit: {git_status}")
        return True
        
    except Exception as e:
        print(f"[ERROR] Error updating documentation: {e}")
        return False


def main():
    """Main entry point."""
    print("\n" + "="*60)
    print("[EN] Updating English Documentation")
    print("="*60)
    
    success = update_documentation()
    
    if success:
        print("\n[OK] Documentation update completed successfully!")
    else:
        print("\n[ERROR] Documentation update failed!")
    
    print("="*60 + "\n")


if __name__ == "__main__":
    main()
