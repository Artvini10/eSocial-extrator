#!/usr/bin/env python3
"""
Gerar documentação completa em Word para o projeto eSocial Extractor (Português)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_style(doc, text, level=1):
    """Adicionar título estilizado"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def shade_cell(cell, color):
    """Colorir célula da tabela"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Criar documento
doc = Document()

# Configurar fonte padrão
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ==================== PÁGINA DE TÍTULO ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('Extrator eSocial\n')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Documentação Completa do Projeto\n')
subtitle_run.font.size = Pt(16)
subtitle_run.font.italic = True

doc_info = doc.add_paragraph()
doc_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc_info_run = doc_info.add_run('Gerado: 29 de Janeiro de 2026')
doc_info_run.font.size = Pt(11)

doc.add_paragraph()  # Espaçamento

# ==================== ÍNDICE ====================
add_heading_style(doc, 'Índice', level=1)
toc_items = [
    '1. Visão Geral do Projeto',
    '2. Arquitetura e Estrutura',
    '3. Estrutura de Diretórios',
    '4. Módulos Principais Explicados',
    '5. Funcionalidades-Chave',
    '6. Como a Aplicação Funciona',
    '7. Configuração e Instalação',
    '8. Dependências',
    '9. Testes e Verificação',
    '10. Guia de Implantação',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ==================== 1. VISÃO GERAL ====================
add_heading_style(doc, '1. Visão Geral do Projeto', level=1)

doc.add_heading('O que é o Extrator eSocial?', level=2)
doc.add_paragraph(
    'O Extrator eSocial é uma aplicação web baseada em FastAPI que orquestra a extração, '
    'transformação e submissão de dados de RH/Folha de Pagamento brasileiros para a plataforma eSocial. '
    'Integra RPA (automação de navegador via Playwright), serviços web governamentais e pipelines de '
    'transformação de dados para automatizar a submissão de documentos e eventos trabalhistas.'
)

doc.add_heading('Propósito', level=2)
doc.add_paragraph(
    'A aplicação funciona como uma ponte entre sistemas internos de RH e a plataforma eSocial brasileira, '
    'permitindo que empresas:'
)
purposes = [
    'Extrair automaticamente dados de RH de diversas fontes',
    'Mapear eventos de RH para códigos de eventos eSocial correspondentes',
    'Submeter dados para APIs do eSocial com assinaturas digitais',
    'Rastrear o status de submissão e tratar respostas',
    'Manter trilhas de auditoria e versionamento de submissões',
]
for purpose in purposes:
    doc.add_paragraph(purpose, style='List Bullet')

doc.add_heading('Tecnologias Principais', level=2)
tech_table = doc.add_table(rows=1, cols=2)
tech_table.style = 'Light Grid Accent 1'
hdr_cells = tech_table.rows[0].cells
hdr_cells[0].text = 'Tecnologia'
hdr_cells[1].text = 'Propósito'

technologies = [
    ('FastAPI', 'Framework web para API REST e templating HTML'),
    ('Playwright', 'Automação RPA para interações com navegador'),
    ('Requests/httpx', 'Cliente HTTP para integração de serviços web'),
    ('Pandas', 'Manipulação e transformação de dados'),
    ('LXML', 'Análise e manipulação de XML'),
    ('signxml', 'Assinaturas digitais XML (requisito eSocial)'),
    ('cryptography', 'Gerenciamento de certificados e chaves'),
    ('Structlog', 'Logging estruturado'),
    ('PyYAML', 'Gerenciamento de arquivos de configuração'),
]
for tech, purpose in technologies:
    row_cells = tech_table.add_row().cells
    row_cells[0].text = tech
    row_cells[1].text = purpose

doc.add_page_break()

# ==================== 2. ARQUITETURA ====================
add_heading_style(doc, '2. Arquitetura e Estrutura', level=1)

doc.add_heading('Fluxo de Dados em Três Camadas', level=2)
doc.add_paragraph(
    'A aplicação segue uma arquitetura em três camadas para processamento de dados:'
)

# Criar fluxo de arquitetura
arch_table = doc.add_table(rows=4, cols=3)
arch_table.style = 'Light Grid Accent 1'

# Cabeçalho
header_cells = arch_table.rows[0].cells
for i, text in enumerate(['Camada', 'Componente', 'Responsabilidade']):
    header_cells[i].text = text
    shade_cell(header_cells[i], 'D3D3D3')

# Dados
arch_data = [
    ('1', 'Adaptador RPA', 'Extrai dados de gov.br via automação de navegador Playwright'),
    ('2', 'Orquestrador', 'Mapeia temas para códigos de eventos e coordena fluxo de trabalho'),
    ('3', 'Adaptador WS', 'Submete para API eSocial e rastreia resultados'),
]
for row_num, (tier, component, resp) in enumerate(arch_data, 1):
    cells = arch_table.rows[row_num].cells
    cells[0].text = tier
    cells[1].text = component
    cells[2].text = resp

doc.add_heading('Pipeline de Processamento de Dados', level=2)
pipeline_steps = [
    'Extração: adaptador RPA recupera dados de RH de gov.br',
    'Transformação: orquestrador mapeia dados para eventos eSocial',
    'Carregamento: escritor exporta para Excel com versionamento',
    'Submissão: adaptador WS envia para API eSocial',
    'Rastreamento: Simulador/API real fornece status de submissão',
]
for i, step in enumerate(pipeline_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_page_break()

# ==================== 3. ESTRUTURA DE DIRETÓRIOS ====================
add_heading_style(doc, '3. Estrutura de Diretórios', level=1)

structure_text = '''eSocial-extrator/
├── src/                              # Código-fonte principal
│   ├── main.py                       # Ponto de entrada da aplicação FastAPI
│   ├── __init__.py                   # Marcador de pacote Python
│   │
│   ├── services/                     # Camada de lógica de negócio
│   │   ├── orchestrator.py           # Mapeamento de tema para código de evento
│   │   ├── ws_adapter.py             # Integração de serviços web
│   │   ├── ws_simulador.py           # Mock de WS para testes
│   │   ├── ws_config.py              # Configuração de WS
│   │   ├── lineage.py                # Versionamento e rastreamento de eventos
│   │   ├── rpa_adapter.py            # Automação RPA/Playwright
│   │   ├── mtls_client.py            # Manipulação de certificados mTLS
│   │   ├── xml_signer.py             # Assinaturas digitais XML
│   │   ├── cert_manager.py           # Carregamento de certificados
│   │   └── __init__.py               # Marcador de pacote
│   │
│   ├── routes/                       # Manipuladores de rotas web
│   │   ├── onboarding.py             # Rotas principais da aplicação
│   │   └── __init__.py               # Marcador de pacote
│   │
│   ├── templates/                    # Templates HTML
│   │   ├── acesso.html               # Formulário de login/acesso
│   │   └── checklist.html            # Página do fluxo principal
│   │
│   ├── static/                       # Ativos estáticos
│   │   ├── logo.png                  # Logo da empresa
│   │   └── style.css                 # Folha de estilos
│   │
│   ├── common/                       # Utilitários compartilhados
│   │   ├── errors.py                 # Classes de exceção personalizadas
│   │   ├── logging_setup.py          # Configuração de logging (pronto)
│   │   ├── ai_helper.py              # Integração com IA (pronto)
│   │   └── __init__.py               # Marcador de pacote
│   │
│   ├── load/                         # Camada de exportação de dados
│   │   ├── writer.py                 # Funcionalidade de exportação Excel
│   │   └── __init__.py               # Marcador de pacote
│   │
│   └── transform/                    # Transformação de dados (pronto para extensão)
│       └── __init__.py               # Marcador de pacote
│
├── config/                           # Arquivos de configuração
│   ├── events_supported.yaml         # Mapeamentos de tema para evento eSocial
│   └── clients.yaml                  # Credenciais de cliente (vazio/template)
│
├── .github/                          # Integração com GitHub
│   ├── copilot-instructions.md       # Diretrizes do assistente de IA
│   └── workflows/                    # Workflows de CI/CD (pronto)
│
├── .env.example                      # Template de variáveis de ambiente
├── .gitignore                        # Regras de ignore do Git
├── requirements.txt                  # Dependências Python
├── test_project.py                   # Suite de testes completa do projeto
├── test_fastapi.py                   # Testes de rotas FastAPI
├── test_ws.py                        # Testes de integração de serviços web
│
├── BUG_FIXES_REPORT.md               # Documentação de correções de bugs
├── VERIFICATION_REPORT.md            # Resultados de testes e verificação
├── FINAL_STATUS.md                   # Status e próximos passos
└── README.md                         # README do projeto'''

doc.add_paragraph(structure_text, style='Normal')

doc.add_page_break()

# ==================== 4. MÓDULOS PRINCIPAIS ====================
add_heading_style(doc, '4. Módulos Principais Explicados', level=1)

modules_info = {
    'src/main.py': {
        'purpose': 'Ponto de entrada da aplicação',
        'functions': [
            'Inicializar aplicação FastAPI',
            'Montar arquivos estáticos',
            'Incluir manipuladores de rotas',
            'Definir endpoint POST /executar_ws',
        ]
    },
    'src/services/orchestrator.py': {
        'purpose': 'Mapeamento de tema para código de evento',
        'functions': [
            'carregar_eventos(): Carregar configuração YAML',
            'expandir_temas(): Converter temas em códigos de evento',
            'Tratamento de erros para temas inválidos',
        ]
    },
    'src/services/ws_adapter.py': {
        'purpose': 'Integração de serviços web',
        'functions': [
            'Classe WSEsocialAdapter',
            'executar(): Execução do fluxo principal',
            'ping_wsdl(): Verificar disponibilidade de WS',
            'Mecanismo de polling para resultados assíncronos',
        ]
    },
    'src/services/ws_simulador.py': {
        'purpose': 'Mock de WS para testes',
        'functions': [
            'Classe WSEsocialSimulador',
            'enviar_lote(): Simular submissão',
            'consultar_lote(): Simular verificação de status',
        ]
    },
    'src/services/lineage.py': {
        'purpose': 'Versionamento e rastreamento de eventos',
        'functions': [
            'aplicar_historico(): Adicionar rastreamento de versão',
            'Marca versão mais recente com flag isAtual',
            'Rastreia histórico de submissão',
        ]
    },
    'src/routes/onboarding.py': {
        'purpose': 'Manipuladores de rotas web',
        'functions': [
            'acesso(): GET / - Formulário de acesso',
            'salvar_acesso(): POST /acesso - Tratar login',
            'checklist(): GET /checklist - Fluxo principal',
            'executar_checklist(): POST /checklist - Processar submissão',
        ]
    },
    'src/common/errors.py': {
        'purpose': 'Tratamento de exceção personalizado',
        'functions': [
            'Classe ErroExplicado com campos causa e solucao',
            'Mensagens de erro melhoradas para usuários',
        ]
    },
    'src/load/writer.py': {
        'purpose': 'Exportação de dados para Excel',
        'functions': [
            'escrever(): Exportar DataFrame para Excel',
            'Cria arquivos *_ATUAL.xlsx e *_HISTORICO.xlsx',
            'Divide dados atuais e históricos',
        ]
    },
}

for module, info in modules_info.items():
    doc.add_heading(module, level=2)
    doc.add_paragraph(f"Propósito: {info['purpose']}", style='Normal')
    doc.add_paragraph("Funções Principais:", style='Normal')
    for func in info['functions']:
        doc.add_paragraph(func, style='List Bullet')

doc.add_page_break()

# ==================== 5. FUNCIONALIDADES-CHAVE ====================
add_heading_style(doc, '5. Funcionalidades-Chave', level=1)

features = {
    'Mapeamento de Eventos Baseado em Tema': {
        'description': 'Mapear temas de RH para códigos de evento eSocial via configuração YAML',
        'example': 'Tema "Vinculos" → Eventos [S-2190, S-2200, S-2205, ...]'
    },
    'Integração de Serviços Web': {
        'description': 'Submeter eventos para API eSocial e rastrear status de submissão',
        'example': 'Simulador mock para testes, suporte de API real com mTLS'
    },
    'Assinaturas Digitais': {
        'description': 'Assinar payloads XML com certificados (requisito eSocial)',
        'example': 'Usa biblioteca signxml com certificados cryptography'
    },
    'Versionamento de Eventos': {
        'description': 'Rastrear múltiplas versões de eventos com divisão atual/histórico',
        'example': 'Flag isAtual marca versão mais recente, outras no histórico'
    },
    'Gerenciamento de Certificados': {
        'description': 'Manipular certificados PFX/PKCS12 e autenticação mTLS',
        'example': 'Extrair chave/cert de PFX, criar arquivos PEM temporários'
    },
    'Exportação de Dados': {
        'description': 'Exportar dados processados para Excel com versionamento',
        'example': 'Arquivos separados *_ATUAL.xlsx e *_HISTORICO.xlsx'
    },
}

for feature, details in features.items():
    doc.add_heading(feature, level=2)
    doc.add_paragraph(f"Descrição: {details['description']}")
    doc.add_paragraph(f"Exemplo: {details['example']}")

doc.add_page_break()

# ==================== 6. COMO FUNCIONA ====================
add_heading_style(doc, '6. Como a Aplicação Funciona', level=1)

doc.add_heading('Fluxo Típico do Usuário', level=2)

workflow_steps = [
    ('Acesso do Usuário', 'Usuário navega para http://localhost:8000'),
    ('Formulário de Login', 'Preenche CNPJ, ambiente, tipo de certificado'),
    ('Submissão de Formulário', 'POST /acesso valida e salva sessão'),
    ('Página de Checklist', 'Usuário seleciona temas e intervalo de datas'),
    ('Expansão de Tema', 'Orquestrador converte temas em códigos de evento'),
    ('Submissão WS', 'Adaptador WS envia eventos para eSocial (ou simulador)'),
    ('Polling', 'App faz poll de resultados até conclusão'),
    ('Exibição de Resultados', 'Mostra status de submissão e recibos'),
    ('Exportação de Dados', 'Resultados exportados para Excel'),
]

workflow_table = doc.add_table(rows=1, cols=2)
workflow_table.style = 'Light Grid Accent 1'
hdr_cells = workflow_table.rows[0].cells
hdr_cells[0].text = 'Passo'
hdr_cells[1].text = 'Descrição'
for step, description in workflow_steps:
    row_cells = workflow_table.add_row().cells
    row_cells[0].text = step
    row_cells[1].text = description

doc.add_heading('Fluxo Técnico', level=2)
doc.add_paragraph('1. Navegador envia POST /acesso com CNPJ e ambiente')
doc.add_paragraph('2. Manipulador de rota cria cookies de sessão')
doc.add_paragraph('3. Usuário submete checklist com temas selecionados')
doc.add_paragraph('4. expandir_temas() converte temas em códigos de evento')
doc.add_paragraph('5. WSEsocialAdapter.executar() inicia submissão')
doc.add_paragraph('6. enviar_lote() envia para API eSocial (ou simulador)')
doc.add_paragraph('7. Recebe número de protocolo e tempo de processamento estimado')
doc.add_paragraph('8. Aguarda tempo estimado e então faz poll de resultados')
doc.add_paragraph('9. Retorna resultados com status de submissão')

doc.add_page_break()

# ==================== 7. CONFIGURAÇÃO ====================
add_heading_style(doc, '7. Configuração e Instalação', level=1)

doc.add_heading('Variáveis de Ambiente (.env)', level=2)
env_vars = [
    ('CERT_PATH', 'Caminho para arquivo de certificado PFX'),
    ('CERT_PASSWORD', 'Senha do certificado'),
    ('CA_BUNDLE_PATH', 'Bundle de CA opcional para verificação SSL'),
    ('SIMULACAO', 'true/false - Usar API mock ou real'),
    ('ESOCIAL_ENVIO_URL_PRODUCAO', 'Endpoint de submissão eSocial real'),
    ('ESOCIAL_CONSULTA_URL_PRODUCAO', 'Endpoint de consulta eSocial real'),
    ('ESOCIAL_ENVIO_URL_PRODUCAO_RESTRITA', 'Submissão de modo restrito'),
    ('ESOCIAL_CONSULTA_URL_PRODUCAO_RESTRITA', 'Consulta de modo restrito'),
]

env_table = doc.add_table(rows=1, cols=2)
env_table.style = 'Light Grid Accent 1'
hdr_cells = env_table.rows[0].cells
hdr_cells[0].text = 'Variável'
hdr_cells[1].text = 'Descrição'
for var, desc in env_vars:
    row_cells = env_table.add_row().cells
    row_cells[0].text = var
    row_cells[1].text = desc

doc.add_heading('events_supported.yaml', level=2)
doc.add_paragraph('Mapeia temas de RH para códigos de evento eSocial. Estrutura:')
yaml_example = '''Vinculos:
  - S-2190
  - S-2200
  - S-2205
  
Dependentes:
  - S-2200
  - S-2205
  - S-1210'''
doc.add_paragraph(yaml_example, style='Normal')

doc.add_page_break()

# ==================== 8. DEPENDÊNCIAS ====================
add_heading_style(doc, '8. Dependências', level=1)

doc.add_heading('Pacotes Python', level=2)
doc.add_paragraph('Veja requirements.txt para lista completa. Dependências principais:')

dep_table = doc.add_table(rows=1, cols=2)
dep_table.style = 'Light Grid Accent 1'
hdr_cells = dep_table.rows[0].cells
hdr_cells[0].text = 'Pacote'
hdr_cells[1].text = 'Versão/Propósito'

dependencies = [
    ('fastapi', 'Framework web'),
    ('uvicorn', 'Servidor ASGI'),
    ('jinja2', 'Templating HTML'),
    ('pandas', 'Manipulação de dados'),
    ('openpyxl', 'Geração de Excel'),
    ('requests/httpx', 'Cliente HTTP'),
    ('playwright', 'Automação RPA'),
    ('lxml', 'Processamento de XML'),
    ('signxml', 'Assinaturas digitais'),
    ('cryptography', 'Manipulação de certificados'),
    ('pyyaml', 'Análise de YAML'),
    ('structlog', 'Logging estruturado'),
]

for package, purpose in dependencies:
    row_cells = dep_table.add_row().cells
    row_cells[0].text = package
    row_cells[1].text = purpose

doc.add_page_break()

# ==================== 9. TESTES ====================
add_heading_style(doc, '9. Testes e Verificação', level=1)

doc.add_heading('Suites de Testes', level=2)

test_suites = {
    'test_project.py': 'Validação completa de estrutura do projeto - 40+ testes',
    'test_fastapi.py': 'Testes de rotas e renderização de templates FastAPI',
    'test_ws.py': 'Testes de funcionalidade de orquestrador e serviços web',
}

for test_file, description in test_suites.items():
    doc.add_paragraph(f'{test_file}: {description}', style='List Bullet')

doc.add_heading('Executando Testes', level=2)
doc.add_paragraph('python test_project.py', style='Normal')
doc.add_paragraph('python test_fastapi.py', style='Normal')
doc.add_paragraph('python test_ws.py', style='Normal')

doc.add_heading('Resultados de Testes', level=2)
doc.add_paragraph('✓ Todos os imports funcionando (12 módulos testados)')
doc.add_paragraph('✓ Todas as rotas funcionais (3 endpoints)')
doc.add_paragraph('✓ Funções principais funcionando (8 funções testadas)')
doc.add_paragraph('✓ Serviços web integrados (simulador + adaptador)')
doc.add_paragraph('✓ Tratamento de erros implementado e testado')
doc.add_paragraph('✓ Zero erros ou avisos encontrados')

doc.add_page_break()

# ==================== 10. IMPLANTAÇÃO ====================
add_heading_style(doc, '10. Guia de Implantação', level=1)

doc.add_heading('Configuração de Desenvolvimento', level=2)
dev_steps = [
    'Clonar repositório',
    'Criar ambiente virtual Python',
    'Instalar dependências: pip install -r requirements.txt',
    'Criar arquivo .env com configuração',
    'Executar servidor de desenvolvimento: python -m uvicorn src.main:app --reload',
]
for i, step in enumerate(dev_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('Implantação em Produção', level=2)
prod_steps = [
    'Configurar endpoints reais da API eSocial em .env',
    'Instalar certificados ICP-Brasil válidos',
    'Definir SIMULACAO=false',
    'Implantar com servidor ASGI de produção (Gunicorn + Uvicorn)',
    'Configurar HTTPS/SSL',
    'Configurar monitoramento e logging',
    'Configurar banco de dados para persistência',
]
for i, step in enumerate(prod_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('Exemplo de Comando de Implantação', level=2)
doc.add_paragraph(
    'gunicorn -w 4 -k uvicorn.workers.UvicornWorker src.main:app --bind 0.0.0.0:8000',
    style='Normal'
)

doc.add_page_break()

# ==================== APÊNDICE ====================
add_heading_style(doc, 'Apêndice: Notas Importantes', level=1)

doc.add_heading('Considerações de Segurança', level=2)
security_notes = [
    'Nunca fazer commit do arquivo .env ou certificados no controle de versão',
    'Usar .gitignore para excluir arquivos sensíveis',
    'Certificados são armazenados em arquivos temporários e limpos',
    'Todas as solicitações HTTP para eSocial usam autenticação mTLS',
]
for note in security_notes:
    doc.add_paragraph(note, style='List Bullet')

doc.add_heading('Aprimoramentos Futuros', level=2)
enhancements = [
    'Implementar camada de transformação para mapeamentos customizados',
    'Adicionar camada de banco de dados para persistência de resultados',
    'Configurar logging estruturado com structlog',
    'Adicionar transformações assistidas por LLM',
    'Criar workflows de CI/CD com GitHub Actions',
    'Adicionar cobertura de testes completa com pytest',
]
for enhancement in enhancements:
    doc.add_paragraph(enhancement, style='List Bullet')

doc.add_heading('Suporte e Recursos', level=2)
doc.add_paragraph('GitHub: https://github.com/Artvini10/eSocial-extrator')
doc.add_paragraph('Documentação: Veja .github/copilot-instructions.md')
doc.add_paragraph('Relatórios de Testes: Veja VERIFICATION_REPORT.md')
doc.add_paragraph('Correções de Bugs: Veja BUG_FIXES_REPORT.md')

# Salvar documento
doc.save('eSocial_Extractor_Documentacao_PTBR.docx')
print("✓ Documento criado: eSocial_Extractor_Documentacao_PTBR.docx")
